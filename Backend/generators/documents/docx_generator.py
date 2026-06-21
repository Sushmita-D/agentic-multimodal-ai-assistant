import os

from docx import Document

OUTPUT_FOLDER = "generators/outputs"

os.makedirs(
    OUTPUT_FOLDER,
    exist_ok=True
)


def generate_docx(
    title: str,
    content: str
):

    file_path = os.path.join(
        OUTPUT_FOLDER,
        "Study_Notes.docx"
    )

    document = Document()

    document.add_heading(
        title,
        level=1
    )

    document.add_paragraph(
        content
    )

    document.save(
        file_path
    )

    return file_path